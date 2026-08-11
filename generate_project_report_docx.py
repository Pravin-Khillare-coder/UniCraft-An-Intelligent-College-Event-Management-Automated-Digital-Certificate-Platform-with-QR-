import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Primary Navy
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark Slate
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo
    return h

def add_body_paragraph(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_bullet_point(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    r_bold = p.add_run(bold_prefix)
    r_bold.font.name = 'Calibri'
    r_bold.font.size = Pt(11)
    r_bold.font.bold = True
    r_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_callout_box(doc, title, text, bg_color="F1F5F9", border_color="3B82F6"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    
    r_t = p.add_run(f"📌 {title}\n")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(10.5)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    r_b = p.add_run(text)
    r_b.font.name = 'Calibri'
    r_b.font.size = Pt(10)
    r_b.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(6)

def generate_full_project_report(output_filename="Project_Report.docx"):
    doc = docx.Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("UniCraft Academic Project Report  |  Event Management Workflow & Platform Architecture")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("UniCraft Platform Report  •  Department of Computer Science & Engineering")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Cover Page Header
    tag_p = doc.add_paragraph()
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(4)
    tag_run = tag_p.add_run("ACADEMIC PROJECT REPORT")
    tag_run.font.name = 'Calibri'
    tag_run.font.size = Pt(11)
    tag_run.font.bold = True
    tag_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(8)
    title_p.paragraph_format.line_spacing = 1.15
    r_title = title_p.add_run("UniCraft: An Intelligent College Event Management & Automated Digital Certificate Platform with QR Verification")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(16)
    r_sub = sub_p.add_run("Comprehensive Technical Documentation with Deep Focus on Event Management Workflow, Upcoming Campus Events, Admin Posting, Student Registrations, E-Certificates & Real-Time App Notifications")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(11.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Info Table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    info_data = [
        [("Course / Degree:", " Bachelor of Technology (B.Tech) / BCA"), ("Academic Session:", " 2025–2026")],
        [("Department:", " Computer Science & Engineering"), ("Domain:", " Full-Stack Web App Development (MERN / Node)")],
        [("Core Focus Areas:", " Event Lifecycle, Student Registration, E-Certificates & App Notifications"), ("Architecture:", " Client-Server REST API with Dual DB Engine")]
    ]
    for row_idx, row in enumerate(info_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.25)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=70, bottom=70, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl, val = info_data[row_idx][col_idx]
            
            r_lbl = p.add_run(lbl)
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
            r_val = p.add_run(val)
            r_val.font.size = Pt(9.5)
            r_val.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Table of Contents
    add_heading_1(doc, "TABLE OF CONTENTS")
    toc_items = [
        "1. Executive Abstract & System Overview",
        "2. Problem Statement & Strategic Objectives",
        "3. IN-DEPTH EVENT MANAGEMENT WORKFLOW (Core Focus)",
        "   3.1 Phase 1: Admin Event Creation & Publication (Hackathons, Workshops, Fests)",
        "   3.2 Phase 2: Upcoming Events Feed & Registration Deadline Tracking",
        "   3.3 Phase 3: Interactive Student Registration & Seat Allocation",
        "   3.4 Phase 4: Event Execution & Real-Time Attendance Marking",
        "   3.5 Phase 5: Automated Bulk E-Certificate Generation & Instant PDF Download",
        "   3.6 Phase 6: Smartphone QR Code Verification & Institutional Audit",
        "4. REAL-TIME APP NOTIFICATION SYSTEM (System Alerts & Activity Triggers)",
        "   4.1 Notification Architecture & Global AuthContext Bus",
        "   4.2 System Event Triggers (Posting, Registration, Attendance & E-Certificates)",
        "   4.3 Notification UI Drawer & Unread Counter Interface",
        "5. Visual WYSIWYG Certificate Designer Engine",
        "6. System Architecture & Dual-Mode Database Infrastructure",
        "7. Core Data Models & Database Schemas",
        "8. Key Technical Algorithms & Mathematical Models",
        "9. Verification, Quality Assurance & Test Matrix",
        "10. Conclusion & Future Scope Roadmap"
    ]
    for item in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_before = Pt(1)
        p_toc.paragraph_format.space_after = Pt(2)
        r_t = p_toc.add_run(item)
        r_t.font.name = 'Calibri'
        r_t.font.size = Pt(10)
        if item.startswith("3.") or item.startswith("4."):
            r_t.font.bold = True
            r_t.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        else:
            r_t.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- SECTION 1 ---
    add_heading_1(doc, "1. Executive Abstract & System Overview")
    add_body_paragraph(
        doc,
        "In modern academic institutions, managing campus events (such as hackathons, technical symposiums, workshops, and cultural fests), "
        "tracking student registrations, broadcasting upcoming event announcements with registration deadlines, issuing physical paper certificates, "
        "and manually verifying credential authenticity is labor-intensive, costly, error-prone, and vulnerable to document forgery."
    )
    add_body_paragraph(
        doc,
        "UniCraft is an intelligent full-stack web application designed to centralize and automate the complete lifecycle of college events "
        "and academic certifications. Administrators can post event details which are instantly displayed to student users in an interactive 'Upcoming Events' feed. "
        "Students can view event dates, venue locations, category filters, and registration deadlines, registering with a single click. "
        "Upon completion of events (hackathons, workshops, cultural fests), UniCraft automatically generates verified e-certificates downloadable in vector PDF format "
        "embedded with scannable tamper-proof QR codes. Furthermore, an integrated Real-Time In-App Notification System keeps students and administrators continuously informed "
        "about event postings, registration confirmations, attendance verifications, and e-certificate availability."
    )

    # --- SECTION 2 ---
    add_heading_1(doc, "2. Problem Statement & Strategic Objectives")
    add_heading_2(doc, "2.1 Existing System Limitations")
    add_bullet_point(doc, "Delayed Campus Event Announcements: ", "Students miss event participation opportunities due to lack of a centralized dashboard displaying upcoming events and registration deadlines.")
    add_bullet_point(doc, "Manual Registration Bottlenecks: ", "Paper forms or isolated Google Forms cause seat over-subscription, missing records, and lack of real-time seat tracking.")
    add_bullet_point(doc, "Lack of Instant Communication: ", "Students receive no real-time status notifications when an event is updated, attendance is marked, or certificates are issued.")
    add_bullet_point(doc, "Paper Certificate Printing & Fraud: ", "Traditional printed certificates incur graphic design costs, printing delays, and cannot be digitally authenticated by employers.")

    add_heading_2(doc, "2.2 Proposed UniCraft Solution Objectives")
    add_bullet_point(doc, "Admin Event Posting & Immediate Student Visibility: ", "Empower administrators to post event details (title, description, date, time, venue, poster, category, max seats) that publish instantly to student feeds.")
    add_bullet_point(doc, "Upcoming Events Feed with Deadline Indicators: ", "Display upcoming college events sorted chronologically with dynamic seat counter badges and registration availability indicators.")
    add_bullet_point(doc, "One-Click Student Registration: ", "Enable authenticated students to register seamlessly with automated seat limit checks and duplicate registration guards.")
    add_bullet_point(doc, "Automated E-Certificate Issuance & Direct Download: ", "Generate vector PDF e-certificates for verified attendees post-event (hackathons, workshops, fests) with scannable verification QR codes.")
    add_bullet_point(doc, "Real-Time In-App Notification System: ", "Broadcast instant app notifications for event creation, registration confirmations, attendance verification, and e-certificate availability.")

    # --- SECTION 3 ---
    add_heading_1(doc, "3. IN-DEPTH EVENT MANAGEMENT WORKFLOW (Core Focus)")
    add_body_paragraph(
        doc,
        "The core foundation of UniCraft is its structured, end-to-end Event Management Workflow. "
        "The workflow governs how an event transitions from administrative creation to student registration, attendance verification, "
        "e-certificate generation, and verification. The complete lifecycle consists of six interconnected phases:"
    )

    add_callout_box(
        doc,
        "WORKFLOW LIFECYCLE SUMMARY",
        "Phase 1: Admin Event Creation & Publication ➔ Phase 2: Upcoming Events Feed & Deadline Display ➔ "
        "Phase 3: Student Discovery & One-Click Registration ➔ Phase 4: Attendance Verification ➔ "
        "Phase 5: Automated E-Certificate Generation & Vector PDF Download ➔ Phase 6: Mobile QR Verification & Audit Logging"
    )

    # Phase 1
    add_heading_2(doc, "3.1 Phase 1: Admin Event Creation & Publication (Hackathons, Workshops, Fests)")
    add_body_paragraph(
        doc,
        "The event lifecycle starts in the Admin Manage Events Dashboard. Administrators enter full event metadata to schedule campus activities:"
    )
    add_bullet_point(doc, "Event Category Selection: ", "Classifies event as Hackathon, Technical Workshop, Symposium, Tech Talk, or Cultural Fest.")
    add_bullet_point(doc, "Event Details & Schedule: ", "Title, detailed description, event date, start time, physical venue/hall location, organizing department, and poster banner image.")
    add_bullet_point(doc, "Seat Capacity & Rules: ", "Defines maximum seat limits (e.g., 100 seats) to control registration volume.")
    add_bullet_point(doc, "Publication State Toggle: ", "Sets initial status to 'Draft' (private preview) or 'Published' (instantly visible on student portals). Upon publishing, an automated app notification is dispatched to all student users.")

    # Phase 2
    add_heading_2(doc, "3.2 Phase 2: Upcoming Events Feed & Registration Deadline Tracking")
    add_body_paragraph(
        doc,
        "When students log into UniCraft, they are greeted by an interactive 'Upcoming Events' section on the Student Dashboard. "
        "This feed is designed to keep students updated on campus activities:"
    )
    add_bullet_point(doc, "Chronological Sorting: ", "Events are automatically ordered by date, placing upcoming college events and imminent registration deadlines at the top.")
    add_bullet_point(doc, "Category Filtering Badges: ", "Students filter upcoming events by categories such as 'Hackathons', 'Workshops', 'Symposiums', or 'Cultural'.")
    add_bullet_point(doc, "Dynamic Capacity & Deadline Indicators: ", "Real-time indicators display remaining seats (e.g., '15 Seats Left / 100') and registration status badges ('Open', 'Closing Soon', 'Fully Booked').")

    # Phase 3
    add_heading_2(doc, "3.3 Phase 3: Interactive Student Registration & Seat Allocation")
    add_body_paragraph(
        doc,
        "Students click on any published event to view full details and register with one click. The backend executes strict validations:"
    )
    add_bullet_point(doc, "Seat Capacity Guard: ", "Queries active registrations. If registration count equals or exceeds maxSeats, registration is blocked with a 'Fully Booked' message.")
    add_bullet_point(doc, "Duplicate Check: ", "Verifies if the student is already registered for the specific event ID.")
    add_bullet_point(doc, "Registration Confirmation & Alert: ", "Creates a Registration record with status 'Registered' and attendance 'Pending', updates seat counter, and dispatches a notification: 'Successfully registered for [Event Title]!'.")

    # Phase 4
    add_heading_2(doc, "3.4 Phase 4: Event Execution & Real-Time Attendance Marking")
    add_body_paragraph(
        doc,
        "Upon event completion (e.g., hackathon judging finished, workshop completed, festival concluded), "
        "administrators access the Event Registrations Roster to mark attendance:"
    )
    add_bullet_point(doc, "Admin Attendance Control: ", "Admin toggles participant attendance status from 'Pending' to 'Present' or 'Absent'.")
    add_bullet_point(doc, "Attendance Notification Trigger: ", "When attendance is marked 'Present', the app sends an instant notification: 'Your attendance for [Event Title] has been verified as Present!'. Only 'Present' students qualify for e-certificate issuance.")

    # Phase 5
    add_heading_2(doc, "3.5 Phase 5: Automated Bulk E-Certificate Generation & Instant PDF Download")
    add_body_paragraph(
        doc,
        "After attendance is verified, the administrator clicks 'Generate All Certificates'. UniCraft's PDF engine creates high-resolution e-certificates:"
    )
    add_bullet_point(doc, "Sequential Allotment: ", "Assigns unique serial numbers (e.g., CERT-2026-0001) and increments the event counter.")
    add_bullet_point(doc, "PDFKit Vector Rendering: ", "Renders A4 landscape vector PDF certificates featuring college background templates, student name, department, event title, official seals, and signatory blocks.")
    add_bullet_point(doc, "Embedded Verification QR Code: ", "Embeds a high-density QR code matrix linked to the server's network verification URL.")
    add_bullet_point(doc, "Student Download Portal & Alert: ", "Certificates immediately appear on the student's 'My Certificates' page with a one-click 'Download PDF' button, and an app notification is sent: 'New certificate issued for [Event Title]! Download it now.'.")

    # Phase 6
    add_heading_2(doc, "3.6 Phase 6: Smartphone QR Code Verification & Institutional Audit")
    add_body_paragraph(
        doc,
        "Anyone scanning the QR code on a printed or digital e-certificate with a smartphone camera is instantly redirected to the public verification portal. "
        "The portal validates the cryptographic verification code and displays verified student credentials, event details, issue timestamp, and official PDF download."
    )

    # --- SECTION 4 ---
    add_heading_1(doc, "4. REAL-TIME APP NOTIFICATION SYSTEM (System Alerts & Activity Triggers)")
    add_body_paragraph(
        doc,
        "UniCraft includes a centralized Real-Time In-App Notification System to keep both students and administrators updated throughout the event lifecycle. "
        "The system is powered by a global React context bus (AuthContext) and persistent storage."
    )

    add_heading_2(doc, "4.1 Notification Architecture & Global AuthContext Bus")
    add_body_paragraph(
        doc,
        "The notification state is maintained in AuthContext.jsx via the addNotification(message) dispatcher. "
        "Any component across the application (Admin Dashboards, Student Pages, Certificate Engine) can emit real-time system alerts."
    )

    add_heading_2(doc, "4.2 Notification System Event Triggers")
    
    ntf_table = doc.add_table(rows=5, cols=3)
    ntf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ntf_table.autofit = False
    set_table_borders(ntf_table, color="CBD5E1")

    n_headers = ["Event Trigger", "Target Recipient", "Sample App Notification Message"]
    n_widths = [Inches(1.8), Inches(1.5), Inches(3.2)]

    for i, ntext in enumerate(n_headers):
        ntf_table.rows[0].cells[i].width = n_widths[i]
        set_cell_background(ntf_table.rows[0].cells[i], "1E3A8A")
        set_cell_margins(ntf_table.rows[0].cells[i], top=80, bottom=80, left=100, right=100)
        p = ntf_table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(ntext)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    ntf_data = [
        ("Admin Posts / Updates Event", "All Students", "📢 'Admin posted new event: CodeSprint 2.0 (Hackathon). Register before seats fill up!'"),
        ("Student Registers for Event", "Student User", "✅ 'Successfully registered for AI/ML Workshop! Your seat is confirmed.'"),
        ("Admin Marks Attendance", "Student User", "🎓 'Your attendance for AI/ML Workshop has been verified as Present!'"),
        ("E-Certificate Issued", "Student User", "📜 'New certificate issued for AI/ML Workshop (CERT-2026-0001). Download it now!'")
    ]

    for r_idx, n_tuple in enumerate(ntf_data):
        row_cells = ntf_table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(n_tuple):
            row_cells[c_idx].width = n_widths[c_idx]
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_heading_2(doc, "4.3 Notification UI Drawer & Unread Counter")
    add_bullet_point(doc, "Header Notification Bell Badge: ", "Displays a dynamic unread count pill badge (e.g., '2 unread') in the top navigation bar.")
    add_bullet_point(doc, "Interactive Dropdown Drawer: ", "Clicking the bell icon slides down a clean notifications panel showing timestamps ('Just now', '2 hours ago'), message details, and 'Mark all as read' controls.")

    # --- SECTION 5 ---
    add_heading_1(doc, "5. Visual WYSIWYG Certificate Designer Engine")
    add_body_paragraph(
        doc,
        "UniCraft includes an interactive browser-based drag-and-drop certificate designer. "
        "Administrators visually position elements on a 640x452 px canvas previewing the final landscape A4 PDF layout."
    )
    add_bullet_point(doc, "Draggable Text Elements: ", "Student Name, Event Title, Date, Serial Number, Issuer Line positioned via mouse drag.")
    add_bullet_point(doc, "Typography & Colors: ", "Font sizes, bold styles, text color pickers, and alignment controls.")
    add_bullet_point(doc, "Custom Stamp & Seal Overlay: ", "Upload college stamp image assets and adjust canvas X/Y placement.")
    add_bullet_point(doc, "Multi-Signatory Cards: ", "Add up to 3 signature cards (Principal, HOD, Coordinator) with live base64 signature image preview.")

    # --- SECTION 6 ---
    add_heading_1(doc, "6. System Architecture & Dual-Mode Database Infrastructure")
    add_body_paragraph(
        doc,
        "UniCraft utilizes a modular Client-Server RESTful Architecture with a Dual-Mode Database Adapter:"
    )
    add_bullet_point(doc, "MongoDB (Mongoose ORM): ", "Primary database for cloud production deployments.")
    add_bullet_point(doc, "Local JSON Database Fallback: ", "Automatic fallback initializing local JSON database files in backend/data/ for offline testing without MongoDB configuration.")

    # --- SECTION 7 ---
    add_heading_1(doc, "7. Core Data Models & Database Schemas")
    add_heading_2(doc, "7.1 User Model Schema")
    add_callout_box(
        doc,
        "USER SCHEMA DEFINITION",
        "_id: ObjectId (Primary Key)\n"
        "name: String (Required)\n"
        "email: String (Required, Unique)\n"
        "password: String (Hashed JWT Credential)\n"
        "role: String ('admin' | 'student')\n"
        "profile: { rollNumber: String, department: String, avatar: String }"
    )

    add_heading_2(doc, "7.2 Event Model Schema")
    add_callout_box(
        doc,
        "EVENT SCHEMA DEFINITION",
        "_id: ObjectId (Primary Key)\n"
        "title: String | description: String | date: Date | time: String | venue: String\n"
        "category: String ('Hackathons', 'Workshops', 'Cultural') | organizer: String | maxSeats: Number\n"
        "status: 'Draft' | 'Published'\n"
        "nextCertificateNumber: String (e.g., 'CERT-2026-0001')\n"
        "certificateTemplate: Base64 String | certificateLayout: Object\n"
        "signatures: Array of { name: String, title: String, signatureImage: Base64 }"
    )

    add_heading_2(doc, "7.3 Certificate Model Schema")
    add_callout_box(
        doc,
        "CERTIFICATE SCHEMA DEFINITION",
        "_id: ObjectId (Primary Key)\n"
        "certificateId: String (Unique Serial Number, e.g., CERT-2026-0001)\n"
        "studentId: ObjectId (Ref: User)\n"
        "eventId: ObjectId (Ref: Event)\n"
        "verificationCode: String (Cryptographic Hash, e.g., e7f8a9b2c3d4)\n"
        "issuedAt: Date (ISO Timestamp)"
    )

    # --- SECTION 8 ---
    add_heading_1(doc, "8. Key Technical Algorithms & Mathematical Models")
    add_heading_2(doc, "8.1 Canvas-to-PDF Coordinate Scaling Algorithm")
    add_body_paragraph(
        doc,
        "To achieve 100% visual parity between the 640x452 px browser canvas and the 842x595 pt PDFKit A4 landscape page:\n"
        "Scale_X = 640 / 842 ≈ 0.76009 \n"
        "Scale_Y = 452 / 595 ≈ 0.75966 \n"
        "PDF_X = Canvas_X / Scale_X \n"
        "PDF_Y = Canvas_Y / Scale_Y"
    )

    add_heading_2(doc, "8.2 Dynamic LAN IP Network Resolution Algorithm")
    add_body_paragraph(
        doc,
        "Inspects OS network interfaces to dynamically bind server Wi-Fi IP for scannable mobile QR verification URLs."
    )

    # --- SECTION 9 ---
    add_heading_1(doc, "9. Verification, Quality Assurance & Test Matrix")
    test_table = doc.add_table(rows=7, cols=4)
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_table.autofit = False
    set_table_borders(test_table, color="CBD5E1")

    t_headers = ["Test ID", "Workflow Step / Feature", "Expected System Behavior", "Result"]
    t_widths = [Inches(0.9), Inches(2.1), Inches(2.5), Inches(1.0)]

    for i, thtext in enumerate(t_headers):
        test_table.rows[0].cells[i].width = t_widths[i]
        set_cell_background(test_table.rows[0].cells[i], "1E3A8A")
        set_cell_margins(test_table.rows[0].cells[i], top=80, bottom=80, left=100, right=100)
        p = test_table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(thtext)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    test_cases = [
        ("TC-01", "Admin Event Publication", "Event status 'Published' immediately appears on student feed.", "PASS"),
        ("TC-02", "Upcoming Event Deadline", "Display seats remaining & block registration when maxSeats reached.", "PASS"),
        ("TC-03", "One-Click Student Reg", "Registration creates record & triggers app notification alert.", "PASS"),
        ("TC-04", "E-Certificate Download", "Verified 'Present' students download vector PDF e-certificates.", "PASS"),
        ("TC-05", "App Notification Drawer", "Notification bell updates unread counter badge on system events.", "PASS"),
        ("TC-06", "Mobile QR Code Scan", "QR code opens verification portal displaying verified credentials.", "PASS")
    ]

    for r_idx, t_tuple in enumerate(test_cases):
        row_cells = test_table.rows[r_idx + 1].cells
        bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(t_tuple):
            row_cells[c_idx].width = t_widths[c_idx]
            set_cell_background(row_cells[c_idx], bg)
            set_cell_margins(row_cells[c_idx], top=70, bottom=70, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx == 0 or c_idx == 3:
                r.font.bold = True
                if c_idx == 3:
                    r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # --- SECTION 10 ---
    add_heading_1(doc, "10. Conclusion & Future Scope Roadmap")
    add_heading_2(doc, "10.1 Project Conclusion")
    add_body_paragraph(
        doc,
        "UniCraft modernizes college event management by integrating administrative event posting, upcoming event deadline tracking, "
        "one-click student registration, automated vector PDF e-certificate issuance, QR code authentication, and a real-time app notification system. "
        "The platform eliminates paper waste, prevents credential fraud, and ensures seamless communication between students and college administration."
    )

    add_heading_2(doc, "10.2 Future Scope Roadmap")
    add_bullet_point(doc, "Push Notification Integration: ", "Adding browser Web Push and SMS integration for mobile notifications.")
    add_bullet_point(doc, "Blockchain Verification: ", "Anchoring certificate hash signatures on Ethereum / Polygon for public verification.")
    add_bullet_point(doc, "LinkedIn Credential Sharing: ", "Direct 'Add to LinkedIn' buttons on student certificate portals.")

    # Save with error handling
    try:
        doc.save(output_filename)
        print(f"Project Report Word Document successfully saved: {output_filename}")
    except PermissionError:
        alt_name = output_filename.replace(".docx", "_Updated.docx")
        doc.save(alt_name)
        print(f"File was locked, saved to alternate filename: {alt_name}")

if __name__ == "__main__":
    generate_full_project_report("Project_Report.docx")
    generate_full_project_report("UniCraft_Project_Report.docx")
