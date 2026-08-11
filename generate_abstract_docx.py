import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_abstract_document(output_filename="Abstract.docx"):
    doc = docx.Document()

    # Set page margins (1 inch = 72 pt)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Normal Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # Header
    header = sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("UniCraft Academic Project Abstract  |  Session 2025–2026")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Footer
    footer = sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Page 1 of 1  •  Confidential - For Academic Review Only")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Category / Tagline
    tag_p = doc.add_paragraph()
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(4)
    tag_run = tag_p.add_run("ACADEMIC PROJECT ABSTRACT")
    tag_run.font.name = 'Calibri'
    tag_run.font.size = Pt(10)
    tag_run.font.bold = True
    tag_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo 600

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    title_p.paragraph_format.line_spacing = 1.15
    title_run = title_p.add_run("UniCraft: An Intelligent College Event Management & Automated Digital Certificate Platform with QR Verification")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900

    # Metadata Card (Table)
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False

    meta_data = [
        [("Project Name:", " UniCraft"), ("Domain:", " Full-Stack Web App (MERN/Node)")],
        [("Target Audience:", " Academic Institutions / Colleges"), ("Key Modules:", " Event Management, QR Certificate Engine")]
    ]

    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.25)
            set_cell_background(cell, "F8FAFC") # Slate 50
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl, val = meta_data[row_idx][col_idx]
            
            r_lbl = p.add_run(lbl)
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(9.5)
            r_lbl.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            
            r_val = p.add_run(val)
            r_val.font.size = Pt(9.5)
            r_val.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Spacing after metadata table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 1 Heading
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    r_h1 = h1.add_run("1. Executive Abstract")
    r_h1.font.size = Pt(13)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue

    # Abstract Paragraphs
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.line_spacing = 1.2
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p1.add_run(
        "In modern higher educational institutions, organizing campus events, tracking student participation, issuing physical paper certificates, "
        "and manually verifying credential authenticity is labor-intensive, costly, error-prone, and highly vulnerable to document forgery. "
        "Traditional certificate distribution workflows lack digital authorization standards, making it difficult for employers and academic verifiers "
        "to confirm document legitimacy quickly."
    )

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.2
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r2_bold = p2.add_run("UniCraft ")
    r2_bold.bold = True
    r2_bold.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    r2 = p2.add_run(
        "is a full-stack, enterprise-grade web application engineered to centralize and automate the complete lifecycle of college events and academic certifications. "
        "The platform introduces an interactive, WYSIWYG Drag-and-Drop Visual Certificate Designer. Administrators can visually layout dynamic student attributes "
        "(name, department, achievement title), customize font styling and color palettes, upload institutional background templates, place official college seals, "
        "and position multi-signatory signature blocks (Principal, HOD, Event Coordinators) directly within the browser."
    )

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(12)
    p3.paragraph_format.line_spacing = 1.2
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r3 = p3.add_run(
        "Upon event completion and attendance validation, UniCraft automatically allocates sequential serial numbers (e.g., CERT-2026-0001) and generates high-resolution "
        "vector PDF certificates embedded with Tamper-Proof Verification QR Codes. Scanning the QR code from any smartphone instantly opens a mobile-optimized public verification portal over local network IP protocols, displaying real-time credential status and direct official PDF downloads."
    )

    # Section 2 Heading
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("2. Key Platform Features & Technical Innovations")
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    features = [
        ("End-to-End Event Lifecycle Management: ", "Role-Based Access Control (RBAC) separating Admin event creation, venue allocation, and registration tracking from Student event discovery and participation portals."),
        ("Visual WYSIWYG Certificate Designer: ", "Browser-based drag-and-drop canvas editor supporting pixel-accurate element positioning, font scaling, background image uploads, dynamic multi-signature blocks, and official stamps."),
        ("Automated Bulk PDF Generation & Sequential Allotment: ", "One-click bulk certificate generation engine that calculates serial numbers sequentially and renders high-density vector PDFs using PDFKit."),
        ("Instant Mobile QR Code Verification: ", "Dynamic Local Area Network (LAN) IP resolution generating embedded scannable QR codes that resolve instantly on mobile browser verification pages without hardcoded endpoints."),
        ("Centralized Credential Audit & CSV Export: ", "Administrative audit registry enabling search, filter, and CSV data export of issued credentials for institutional compliance and record keeping.")
    ]

    for title, desc in features:
        fp = doc.add_paragraph(style='List Bullet')
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(4)
        fp.paragraph_format.line_spacing = 1.15
        
        r_t = fp.add_run(title)
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        r_d = fp.add_run(desc)
        r_d.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Section 3 Heading
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    r_h3 = h3.add_run("3. Technology Stack & Architecture")
    r_h3.font.size = Pt(13)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    # Tech Stack Table
    stack_table = doc.add_table(rows=5, cols=3)
    stack_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stack_table.autofit = False
    set_table_borders(stack_table, color="CBD5E1", sz="4")

    headers = ["Layer / Module", "Technologies Used", "Functionality & Purpose"]
    col_widths = [Inches(1.8), Inches(2.2), Inches(2.5)]

    # Header Row
    hdr_cells = stack_table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E3A8A") # Dark Navy
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    stack_data = [
        ("Frontend Client", "React 18, Vite, TailwindCSS, Lucide Icons", "Responsive SPA interface, interactive visual canvas designer, state management."),
        ("Backend REST API", "Node.js, Express.js, JWT Middleware", "RESTful API routes, user authentication, role authorization, business logic."),
        ("PDF & QR Engine", "PDFKit, QRCode Matrix Generator", "Vector PDF rendering, coordinate scaling algorithms, scannable QR code generation."),
        ("Database Layer", "MongoDB (Mongoose) / Local JSON Adapter", "Persistent storage for users, events, registrations, certificates, with offline fallback.")
    ]

    for row_idx, data in enumerate(stack_data):
        row_cells = stack_table.rows[row_idx + 1].cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.font.size = Pt(9)
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            else:
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # Spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 4 Heading
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    r_h4 = h4.add_run("4. Conclusion & Real-World Impact")
    r_h4.font.size = Pt(13)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p_conc = doc.add_paragraph()
    p_conc.paragraph_format.space_before = Pt(0)
    p_conc.paragraph_format.space_after = Pt(12)
    p_conc.paragraph_format.line_spacing = 1.2
    p_conc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_conc.add_run(
        "UniCraft modernizes institutional event workflows by replacing fragmented paper-based processes with an integrated digital ecosystem. "
        "By combining visual template design, automated serial allotment, instant PDF rendering, and smartphone-scannable QR verification, "
        "UniCraft eliminates manual administrative bottlenecks, drastically reduces printing overhead, and guarantees uncompromised certificate authenticity."
    )

    doc.save(output_filename)
    print(f"Document successfully created: {output_filename}")

if __name__ == "__main__":
    build_abstract_document("Abstract.docx")
    build_abstract_document("UniCraft_Abstract.docx")
